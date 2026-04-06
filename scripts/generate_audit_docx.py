import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_audit_docx(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('[AUDIT] Cathedral Megasite Navigation & Flow (v2.0)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_paragraph('This document provides a page-by-page mapping of the site\'s navigation architecture to ensure a seamless "walkthrough" experience.')
    
    # Narrative Mapping Section
    doc.add_heading('Narrative Mapping: Rings vs. Acts', level=1)
    doc.add_paragraph('To reconcile the "Shield" metaphor with the "Investigation" flow, the following mapping is established:')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ring (Shield)'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Narrative Act'
    hdr_cells[3].text = 'Key Dossier Hub'
    
    data = [
        ('Ring 5 (Outer)', 'VICTIMS & SURVIVORS', 'Act IV / VII', '/evidence/victims'),
        ('Ring 4', 'THE OUTER WALLS', 'Act VI / VIII', '/ledger/checchio'),
        ('Ring 3', 'LEGAL APPARATUS', 'Act III', '/ledger/mccarrick/mechanisms'),
        ('Ring 2', 'LEGAL FORTRESS', 'Act II (Labyrinth)', '/ledger/mccarrick/network'),
        ('Ring 1 (Inner)', 'SECRET ARCHIVES', 'Act I (Origin)', '/evidence/origin')
    ]
    
    for ring, name, act, path in data:
        row_cells = table.add_row().cells
        row_cells[0].text = ring
        row_cells[1].text = name
        row_cells[2].text = act
        row_cells[3].text = path
        
    doc.add_paragraph('') # Spacer
    
    # Page-by-Page Audit
    doc.add_heading('Page-by-Page Navigation Audit', level=1)
    
    pages = [
        {
            "title": "1. THE HUB (Home)",
            "details": [
                "Path: /",
                "Tip-top (Global): PrimaryNavigation (Visible)",
                "Top (Context): Hero Narrative + Video Background.",
                "Bottom: HomeNavigation (Act I-IX grid).",
                "Status: Healthy. Entry point for 'The Descent'."
            ]
        },
        {
            "title": "2. THE EVIDENCE (Hub)",
            "details": [
                "Path: /evidence",
                "Tip-top (Global): PrimaryNavigation (Visible)",
                "Top (Context): EvidenceSubNavigation (Archives, Board, Lawyers).",
                "Bottom: JourneyNav (Manual sequence).",
                "Status: WARNING. The 'Begin Act I' CTA should point to Ring 1 (Origin) for narrative consistency."
            ]
        },
        {
            "title": "3. THE LABYRINTH (Act II)",
            "details": [
                "Path: /ledger/mccarrick/network (Primary Mapping)",
                "Tip-top (Global): PrimaryNavigation (Visible)",
                "Top (Context): EndgameLayout Sidebar (McCarrick Dossier).",
                "Bottom: BROKEN. JourneyNav is missing or the path is incorrectly indexed.",
                "Action: Re-integrate JourneyNav and add path to registry."
            ]
        },
        {
            "title": "4. THE BREACH (Portal)",
            "details": [
                "Path: /breach",
                "Tip-top (Global): PrimaryNavigation (Visible)",
                "Top (Context): Forensic Metrics.",
                "Bottom: JourneyNav.",
                "Status: Healthy."
            ]
        },
        {
            "title": "5. THE LEDGER (Investigative Hub)",
            "details": [
                "Path: /ledger",
                "Tip-top (Global): PrimaryNavigation (Visible)",
                "Top (Context): Character selection (McCarrick, Checchio, Lorenzo).",
                "Bottom: JourneyNav.",
                "Status: Healthy."
            ]
        }
    ]
    
    for page in pages:
        doc.add_heading(page['title'], level=2)
        for detail in page['details']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(detail)
            
    # Technical Recommendations
    doc.add_heading('Technical Recommendations', level=1)
    recs = [
        "Unify JourneyNav.tsx: Expand the ROUTES array to include every sub-page in the /ledger/mccarrick and /ledger/checchio dossiers.",
        "Standardize Breadcrumbs: Implement a ForensicBreadcrumbs component for the 'Top' navigation space to show SH > EVIDENCE > RING 2 > LABYRINTH.",
        "Refactor EvidenceHub.tsx: Update the 6-ring grid to map logically to the 5-ring Shield Registry to avoid user confusion."
    ]
    for rec in recs:
        doc.add_paragraph(rec, style='List Number')
        
    doc.save(output_path)

if __name__ == "__main__":
    create_audit_docx(r"c:\Projects\Cathedral_Megasite\sodomhall-project\docs\navigation_audit_report_v2.docx")
